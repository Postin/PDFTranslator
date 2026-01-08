import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_bilingual_docx(
    pages: list,
    output_path: str,
    source_lang: str,
    target_lang: str
):
    """
    Create a Word document with original and translated text for each page.
    Sequential layout: original text first, then translation below.
    
    Args:
        pages: List of dicts with {"original": str, "translated": str, "page_num": int}
        output_path: Output .docx file path
        source_lang: Source language name (for headers)
        target_lang: Target language name (for headers)
    """
    doc = Document()
    
    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    
    for page_data in pages:
        page_num = page_data["page_num"]
        original = page_data.get("original", "")
        translated = page_data.get("translated", "")
        
        # Page header
        header = doc.add_heading(f"Page {page_num}", level=1)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Original text section
        original_header = doc.add_heading(f"Original ({source_lang})", level=2)
        for run in original_header.runs:
            run.font.color.rgb = RGBColor(70, 70, 70)
        
        for para in original.split("\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
            else:
                doc.add_paragraph()
        
        # Separator
        separator = doc.add_paragraph("─" * 60)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Translated text section
        translated_header = doc.add_heading(f"Translation ({target_lang})", level=2)
        for run in translated_header.runs:
            run.font.color.rgb = RGBColor(0, 102, 153)
        
        for para in translated.split("\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
            else:
                doc.add_paragraph()
        
        # Page break after each page
        doc.add_page_break()
    
    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    doc.save(output_path)
    print(f"Bilingual document saved to: {output_path}")
